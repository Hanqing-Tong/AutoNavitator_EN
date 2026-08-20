import os
import io
import base64
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile
from PIL import Image

class ExportService:
    def __init__(self):
        pass

    def _blend_masks_with_original(self, original_base64, targets_masks: list):
        """
        Overlay mask images of multiple targets on the original image with different colors to generate a rendered image
        :param targets_masks: List of target mask Base64 strings
        """
        try:
            def decode_base64(b64_str):
                if not b64_str: return None
                if ',' in b64_str:
                    b64_str = b64_str.split(',')[1]
                return Image.open(io.BytesIO(base64.b64decode(b64_str)))

            orig_img = decode_base64(original_base64)
            if not orig_img:
                return None

            orig_img = orig_img.convert("RGBA")
            
            # Define a high-contrast color palette (RGB), set Alpha value to 100 (approx. 40% opacity) to balance base map visibility and mask recognizability
            color_palette = [
                (255, 0, 0),    # Red
                (0, 255, 0),    # Green
                (0, 0, 255),    # Blue
                (255, 255, 0),  # Yellow
                (255, 0, 255),  # Magenta
                (0, 255, 255),  # Cyan
                (255, 165, 0),  # Orange
                (128, 0, 128),  # Purple
            ]
            alpha_value = 100 # Set back to 100 to increase color intensity
            
            # Create a general overlay layer
            combined_overlay = Image.new("RGBA", orig_img.size, (0, 0, 0, 0))
            
            print(f"Blending {len(targets_masks)} masks...")
            for idx, mask_b64 in enumerate(targets_masks):
                mask_img = decode_base64(mask_b64)
                if not mask_img:
                    print(f"Mask {idx} is empty")
                    continue
                
                mask_img = mask_img.convert("L").resize(orig_img.size)
                
                # Select color for the current target
                color = color_palette[idx % len(color_palette)]
                target_overlay = Image.new("RGBA", orig_img.size, color + (alpha_value,))
                
                # Use the target's mask as the alpha channel
                target_overlay.putalpha(mask_img)
                
                # Merge into the general overlay layer
                combined_overlay = Image.alpha_composite(combined_overlay, target_overlay)

            # Composite the general overlay layer with the original image
            final_img = Image.alpha_composite(orig_img, combined_overlay)
            return final_img.convert("RGB")
        except Exception as e:
            print(f"Blending error: {e}")
            return None

    def _add_latex_image(self, paragraph, latex_code):
        """
        Convert LaTeX formulas to images and insert them into the paragraph
        """
        import requests
        import io
        import urllib.parse
        
        # Remove $ or $$ symbols
        formula = latex_code.strip('$')
        # URL encode the LaTeX code
        encoded_formula = urllib.parse.quote(formula)
        # Use more stable API parameters
        url = f"https://latex.codecogs.com/png.latex?\dpi{{150}} {encoded_formula}"
        
        try:
            # Add User-Agent to avoid being blocked by the API
            headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'}
            response = requests.get(url, headers=headers, timeout=10)
            
            if response.status_code == 200 and len(response.content) > 100: # Ensure the return is an image rather than an error page
                img_stream = io.BytesIO(response.content)
                run = paragraph.add_run()
                # 插入公式图片
                run.add_picture(img_stream)
            else:
                # On failure, remove $ symbols and show plain text to avoid showing LaTeX source code to the user
                paragraph.add_run(formula) 
        except Exception as e:
            print(f"LaTeX render error: {e}")
            paragraph.add_run(formula) # Show plain text on exception

    def _add_formatted_text(self, paragraph, text):
        """
        Parse simple Markdown format (bold, italic, LaTeX formulas) and add them to the paragraph
        """
        import re
        # Match $$block$$, $inline$$, **bold**, *italic*
        # Priority: block formula > inline formula > bold > italic
        pattern = re.compile(r'(\$\$.*?\$\$|\$.*?\$|\*\*.*?\*\*|\*.*?\*)')
        parts = pattern.split(text)
        
        for part in parts:
            if not part:
                continue
            if part.startswith('$$') and part.endswith('$$'):
                self._add_latex_image(paragraph, part)
            elif part.startswith('$') and part.endswith('$'):
                self._add_latex_image(paragraph, part)
            elif part.startswith('**') and part.endswith('**'):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('*') and part.endswith('*'):
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            else:
                paragraph.add_run(part)

    def create_docx_report(self, report_text: str, images_data: list):
        """
        Export the analysis report as a .docx file
        :param report_text: Report text in Markdown format
        :param images_data: List containing image information [{"file_name": "...", "image_base64": "..."}]
        :return: Path to the temporary file
        """
        doc = Document()
        
        # Set title
        title = doc.add_heading('AutoNavitator AI Analysis Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 1. Add imagery evidence section (refer to diagnostic report, placed below the title)
        if images_data:
            doc.add_heading('Imagery Evidence', level=1)
            
            # Calculate number of rows (2 images per row)
            num_imgs = len(images_data)
            num_rows = (num_imgs + 1) // 2
            
            # Create a borderless table for parallel display
            table = doc.add_table(rows=num_rows, cols=2)
            table.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            for idx, img_item in enumerate(images_data):
                row = idx // 2
                col = idx % 2
                cell = table.cell(row, col)
                
                file_name = img_item.get('file_name', 'unknown')
                base64_str = img_item.get('image_base64') or img_item.get('mask_base64', '')
                
                if not base64_str:
                    cell.text = f"Image {file_name} missing"
                    continue
                
                if ',' in base64_str:
                    base64_str = base64_str.split(',')[1]
                
                try:
                    img_bytes = base64.b64decode(base64_str)
                    img_stream = io.BytesIO(img_bytes)
                    
                    # Check if original image and mask image are provided, try to synthesize a rendered image
                    original_b64 = img_item.get('original_base64')
                    mask_b64 = img_item.get('mask_base64')
                    
                    if original_b64:
                        # Use targets list for multi-color rendering, fallback to single mask if not present
                        targets = img_item.get('targets', [])
                        if targets:
                            target_masks = [t.get('mask_base64') for t in targets if t.get('mask_base64')]
                            blended_img = self._blend_masks_with_original(original_b64, target_masks)
                        else:
                            # Fallback: wrap single mask into a list
                            blended_img = self._blend_masks_with_original(original_b64, [mask_b64] if mask_b64 else [])
                        
                        if blended_img:
                            img_byte_arr = io.BytesIO()
                            blended_img.save(img_byte_arr, format='JPEG')
                            img_stream = io.BytesIO(img_byte_arr.getvalue())
                    
                        # Add image and title within the cell
                        # Note: cell.add_paragraph() creates a paragraph by default
                        p_title = cell.paragraphs[0]
                        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = p_title.add_run(f"Imagery File: {file_name}")
                        run.bold = True
                    
                    run_img = cell.add_paragraph().add_run()
                    run_img.add_picture(img_stream, width=Inches(3)) # 宽度约3英寸，确保一行能放两张
                except Exception as e:
                    cell.text = f"Image {file_name} export failed: {str(e)}"
            
            doc.add_paragraph("\n")

        # 2. Add text analysis section
        doc.add_heading('Analysis Conclusion', level=1)
        
        lines = report_text.split('\n')
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            if not line:
                i += 1
                continue
            
            # Detect if it is a Markdown table (starts with | and contains multiple |)
            if line.startswith('|') and line.count('|') >= 2:
                table_lines = []
                while i < len(lines) and lines[i].strip().startswith('|'):
                    table_lines.append(lines[i].strip())
                    i += 1
                
                # Process table data
                # Filter out Markdown separator lines (e.g., |---|---|)
                rows_data = []
                for tl in table_lines:
                    # Remove leading/trailing | and split by |
                    cells = [c.strip() for c in tl.strip('|').split('|')]
                    # Check if it is a separator line
                    if all(set(c).issubset({'-', ':', ' '}) for c in cells):
                        continue
                    rows_data.append(cells)
                
                    if rows_data:
                        # Create Word table
                        table = doc.add_table(rows=len(rows_data), cols=len(rows_data[0]))
                        table.style = 'Table Grid' # Add grid lines
                        for r_idx, row_cells in enumerate(rows_data):
                            for c_idx, cell_text in enumerate(row_cells):
                                if c_idx < len(table.columns):
                                    cell = table.cell(r_idx, c_idx)
                                    # Also perform formatting on cell content (e.g., LaTeX)
                                    p = cell.paragraphs[0]
                                    self._add_formatted_text(p, cell_text)
                continue # 跳过末尾的 i += 1
            
            if line.startswith('###'):
                text = line.replace('###', '').strip()
                heading = doc.add_heading('', level=3)
                self._add_formatted_text(heading, text)
            elif line.startswith('##'):
                text = line.replace('##', '').strip()
                heading = doc.add_heading('', level=2)
                self._add_formatted_text(heading, text)
            elif line.startswith('#'):
                text = line.replace('#', '').strip()
                heading = doc.add_heading('', level=1)
                self._add_formatted_text(heading, text)
            elif line.startswith('- ') or line.startswith('* '):
                text = line[2:].strip()
                p = doc.add_paragraph(style='List Bullet')
                self._add_formatted_text(p, text)
            else:
                p = doc.add_paragraph()
                self._add_formatted_text(p, line)
            i += 1

        # Save to temporary file
        temp_dir = tempfile.gettempdir()
        file_path = os.path.join(temp_dir, "analysis_report.docx")
        doc.save(file_path)
        
        return file_path